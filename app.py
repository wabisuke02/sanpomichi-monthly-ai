import streamlit as st
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import os

load_dotenv()

client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)

st.title("デイサービス散歩道 月次報告AI")

office_name = st.text_input("事業所名", "デイサービス散歩道")
target_month = st.text_input("対象月", "2026年6月")

excel_report = st.text_area(
    "Excelの月次報告書をコピーして、そのまま貼り付けてください",
    height=400,
    placeholder="Excelの表を範囲選択 → コピー → ここに貼り付け"
)

activities = st.text_area(
    "今月やったこと",
    height=150,
    placeholder="例：生活相談員研修、自主トレプリント作成、パンフレット更新 など"
)

result = st.text_area(
    "その結果",
    height=150,
    placeholder="例：利用者の反応、見学につながったこと、職員の変化 など"
)

sales = st.text_area(
    "営業活動",
    height=150,
    placeholder="例：ケアマネ訪問、FAX送信、見学対応、紹介依頼 など"
)
hiyari = st.text_area(
    "ヒヤリ",
    height=120
)

incident = st.text_area(
    "インシデント",
    height=120
)

accident = st.text_area(
    "事故",
    height=120
)

customer_voice = st.text_area(
    "お客様からのご意見",
    height=120,
    placeholder="例：体操の声が大きい、送迎が早い、運動量を増やしたい など"
)

consult = st.text_area(
    "相談事項",
    height=150,
    placeholder="例：新規獲得、稼働率改善、現場課題、営業方法 など"
)

next_action = st.text_area(
    "来月やること",
    height=150,
    placeholder="例：空き情報FAX、ケアマネ営業、計画書研修、見学促進 など"
)

if st.button("報告書作成"):

    prompt = f"""
あなたは介護事業所の管理者会向け報告書作成AIです。

目的：
介護事業所の月次数値表と現場メモから、
管理者会・社内会議で使える月次報告書を作成する。

主な対象：
・デイサービス散歩道
・散歩道金沢

入力される情報：
1. 事業所名
2. 対象月
3. 月次数値表
4. 今月やったこと
5. その結果
6. 営業活動
7. ヒヤリハット・事故
8. 現在困っていること
9. 来月相談したいこと

作成ルール：
作成ルール：
・個人情報は絶対に出さない
・利用者実名は出さない
・職員名も必要がなければ出さない
・数値と事業所単位の情報だけで作成する
・前月比較を行う
・現場メモを反映する
・課題と来月方針を連動させる
・管理者会でそのまま読める文章にする
・ヒヤリ、インシデント、事故は区別して記載する
・入力が空欄の場合は「入力なし」と記載し、「報告なし」と断定しない
・単なる入力内容の要約ではなく、数値と現場情報から分析を行うこと
・課題に対して原因を推測し、具体的な改善案を提案すること
・管理者会で議論できるレベルまで掘り下げること。
・入力内容にない事実は作らない。
・原因分析は「考えられる」「推測される」と明記すること。
・根拠がない断定は禁止。

出力形式：
【月次報告書】
■実績概要

■良かったこと

■課題

■営業活動

■事故・ヒヤリ

■お客様のご意見

■相談事項

■来月方針

事業所名：
{office_name}

対象月：
{target_month}

Excel月次報告書：
{excel_report}

今月やったこと：
{activities}

その結果：
{result}

営業活動：
{sales}

ヒヤリ：
{hiyari}

インシデント：
{incident}

事故：
{accident}

お客様からのご意見：
{customer_voice}

相談事項：
{consult}

来月やること：
{next_action}
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    report_text = response.choices[0].message.content

    st.write(report_text)

    doc = Document()
    doc.add_heading("デイサービス散歩道 月次報告書", level=1)

    for line in report_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    st.download_button(
        label="Wordでダウンロード",
        data=file_stream,
        file_name=f"{target_month}_{office_name}_月次報告書.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if st.button("月次ミーティングレジメ作成"):

    meeting_prompt = f"""
あなたは「デイサービス散歩道 ミーティングレジメ作成AI」です。

目的：
管理者が作成した月次報告書と入力情報から、
「生活相談員がそのまま司会できる月例ミーティングレジメ」
を作成する。

重要事項：
このレジメは報告資料ではない。
司会者（生活相談員）が会議を進行し、
スタッフから意見を引き出し、
話し合いができることを目的とする。

管理者は補足役とし、
会議の主役は生活相談員と現場スタッフとする。

作成ルール：
・司会者が読み上げやすい文章にする
・スタッフへ質問する内容を必ず作る
・意見が出やすい問いかけを作る
・数字報告だけで終わらせない
・改善案が出る会議にする
・A4で2〜3枚程度
・そのまま印刷して使用できる形式
・各項目に【司会コメント】を必ず入れる
・各項目に【みんなに聞く】を必ず入れる
・単なる議事録ではなく司会進行台本として作成する

出力構成：
① 開会
② 委員会・法人共有事項
③ 先月の実績報告
④ 新規・営業報告
⑤ 現場の取り組み
⑥ ヒヤリ・インシデント・事故・お客様のご意見共有
⑦ 現在の課題
⑧ 来月やること
内容｜担当｜期限
⑨ 管理者コメント

入力情報：

事業所名：
{office_name}

対象月：
{target_month}

Excel月次報告書：
{excel_report}

今月やったこと：
{activities}

その結果：
{result}

営業活動：
{sales}

ヒヤリ：
{hiyari}

インシデント：
{incident}

事故：
{accident}

お客様からのご意見：
{customer_voice}

相談事項：
{consult}

来月やること：
{next_action}
"""

    meeting_response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "user", "content": meeting_prompt}
        ]
    )

    meeting_text = meeting_response.choices[0].message.content

    st.write(meeting_text)

    doc = Document()
    doc.add_heading("デイサービス散歩道 月次ミーティングレジメ", level=1)

    for line in meeting_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    st.download_button(
        label="ミーティングレジメをWordでダウンロード",
        data=file_stream,
        file_name=f"{target_month}_{office_name}_月次ミーティングレジメ.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    